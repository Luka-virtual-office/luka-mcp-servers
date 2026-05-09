"""
Word document tools
───────────────────
read_docx   — .docx → structured text (paragraphs, tables, headings)
text_to_docx — Markdown/plain text → .docx (base64)
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import ALLOWED_DOCX_MIMES
from utils.file_utils import encode_b64, load_and_validate, temp_file


# ─────────────────────────────────────────────────────────────────────────────
# read_docx
# ─────────────────────────────────────────────────────────────────────────────

def read_docx(file_content: str) -> dict[str, Any]:
    """
    Extract text and structure from a .docx file (base64-encoded).

    Returns
    -------
    dict with:
        sections     — list of {type, level, text} (paragraphs + headings)
        tables       — list of {headers, rows}
        word_count   — approximate
        char_count   — approximate
        plain_text   — full plain-text concatenation
    """
    raw, _ = load_and_validate(file_content, ALLOWED_DOCX_MIMES, "Word document", ".docx")

    with temp_file(raw, suffix=".docx") as path:
        doc = Document(str(path))

        sections: list[dict] = []
        plain_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"
            heading_match = re.match(r"Heading (\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1))
                sections.append({"type": "heading", "level": level, "text": text})
            else:
                sections.append({"type": "paragraph", "text": text})
            plain_parts.append(text)

        tables_data: list[dict] = []
        for table in doc.tables:
            t_rows = []
            for row in table.rows:
                t_rows.append([cell.text.strip() for cell in row.cells])
            if t_rows:
                headers = t_rows[0]
                data_rows = t_rows[1:]
                tables_data.append({"headers": headers, "rows": data_rows})
                # Also add to plain_text
                plain_parts.append(" | ".join(headers))
                for r in data_rows:
                    plain_parts.append(" | ".join(r))

        plain_text = "\n".join(plain_parts)
        word_count = len(plain_text.split())
        char_count = len(plain_text)

    return {
        "sections": sections,
        "tables": tables_data,
        "word_count": word_count,
        "char_count": char_count,
        "plain_text": plain_text,
        "section_count": len(sections),
        "table_count": len(tables_data),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → docx helpers
# ─────────────────────────────────────────────────────────────────────────────

def _parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """
    Parse inline **bold** and *italic* markers.
    Returns list of (text, is_bold, is_italic).
    """
    parts: list[tuple[str, bool, bool]] = []
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(1).startswith("***"):
            parts.append((m.group(2), True, True))
        elif m.group(1).startswith("**"):
            parts.append((m.group(3), True, False))
        elif m.group(1).startswith("`"):
            parts.append((m.group(5), False, False))
        else:
            parts.append((m.group(4), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts or [(text, False, False)]


def _add_para(doc: Document, text: str, style: str = "Normal") -> None:
    """Add a paragraph with inline formatting parsed from Markdown."""
    p = doc.add_paragraph(style=style)
    for chunk, bold, italic in _parse_inline(text):
        run = p.add_run(chunk)
        run.bold = bold
        run.italic = italic


# ─────────────────────────────────────────────────────────────────────────────
# text_to_docx
# ─────────────────────────────────────────────────────────────────────────────

def text_to_docx(
    content: str,
    title: str | None = None,
    author: str | None = None,
) -> dict[str, str]:
    """
    Convert Markdown (or plain text) to a formatted .docx (base64).

    Supported Markdown:
        # H1  ## H2  ### H3  #### H4
        **bold**  *italic*  ***bold-italic***  `code`
        - bullet list (- or *)
        1. numbered list
        ---  (horizontal rule → page-break style line)
        | table | headers |  (simple pipe tables)
        > blockquote
        ```code block```

    Parameters
    ----------
    content : str
        Markdown or plain text.
    title : str | None
        Document title (written as H1 at the top if provided).
    author : str | None
        Written to core document properties.

    Returns
    -------
    {"file_content": "<base64>", "mime_type": "...", "size_bytes": N}
    """
    doc = Document()

    # Core properties
    if author:
        doc.core_properties.author = author
    if title:
        doc.core_properties.title = title
        doc.add_heading(title, level=0)  # Title style

    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = doc.styles["No Spacing"]
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.get_or_add_pBdr() if hasattr(pPr, "get_or_add_pBdr") else None
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:], style="Intense Quote")
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[-| :]+\|$", lines[i].strip()):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                rows = [
                    [cell.strip() for cell in row.strip().strip("|").split("|")]
                    for row in table_lines
                ]
                num_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = cell_text
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        bullet_match = re.match(r"^(\s*)[*\-]\s+(.+)", line)
        if bullet_match:
            _add_para(doc, bullet_match.group(2), style="List Bullet")
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        num_match = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if num_match:
            _add_para(doc, num_match.group(2), style="List Number")
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        _add_para(doc, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    return {
        "file_content": encode_b64(raw),
        "mime_type": (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        "size_bytes": len(raw),
    }
