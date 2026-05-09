"""
Test suite for MCP Document Server tools.
Run with:  cd mcp-document-server && pytest tests/ -v
"""
from __future__ import annotations

import base64
import io
import json
import sys
from pathlib import Path

import pytest

# Make the parent directory importable
sys.path.insert(0, str(Path(__file__).parent.parent))

# ── Helpers ───────────────────────────────────────────────────────────────────

def b64(data: bytes) -> str:
    return base64.b64encode(data).decode()


def from_b64(s: str) -> bytes:
    return base64.b64decode(s)


# ─────────────────────────────────────────────────────────────────────────────
# Excel
# ─────────────────────────────────────────────────────────────────────────────

class TestDataToExcel:
    """data_to_excel — generate xlsx from JSON data."""

    def test_list_of_dicts(self):
        from tools.excel_tools import data_to_excel
        data = [
            {"name": "Alice", "revenue": 5000, "region": "LATAM"},
            {"name": "Bob",   "revenue": 8200, "region": "US"},
            {"name": "Carol", "revenue": 3100, "region": "EU"},
        ]
        result = data_to_excel(data)
        assert "file_content" in result
        assert result["row_count"] == 3
        assert result["column_count"] == 3
        raw = from_b64(result["file_content"])
        assert raw[:4] == b"PK\x03\x04"  # xlsx magic = zip

    def test_list_of_lists(self):
        from tools.excel_tools import data_to_excel
        data = [
            ["Product", "Q1", "Q2"],
            ["Widget A", 100, 120],
            ["Widget B", 200, 250],
        ]
        result = data_to_excel(data, sheet_name="Sales", include_headers=True)
        assert result["row_count"] == 2
        assert result["column_count"] == 3

    def test_with_title(self):
        from tools.excel_tools import data_to_excel
        data = [{"x": 1, "y": 2}]
        result = data_to_excel(data, title="My Report")
        assert result["row_count"] == 1

    def test_empty_raises(self):
        from tools.excel_tools import data_to_excel
        with pytest.raises(ValueError):
            data_to_excel([])


class TestReadExcel:
    """read_excel — parse xlsx and return structured JSON."""

    def _make_xlsx(self, rows: list[list]) -> str:
        """Create a minimal xlsx in memory and return as base64."""
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "TestSheet"
        for row in rows:
            ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return b64(buf.getvalue())

    def test_basic_read(self):
        from tools.excel_tools import read_excel
        b64_file = self._make_xlsx([
            ["Name", "Score"],
            ["Alice", 95],
            ["Bob", 87],
        ])
        result = read_excel(b64_file)
        assert result["sheet_count"] == 1
        sheet = result["sheets"][0]
        assert sheet["sheet_name"] == "TestSheet"
        assert sheet["headers"] == ["Name", "Score"]
        assert sheet["row_count"] == 2

    def test_specific_sheet(self):
        from tools.excel_tools import read_excel
        b64_file = self._make_xlsx([["A", "B"], [1, 2]])
        result = read_excel(b64_file, sheet_name="TestSheet")
        assert result["sheet_count"] == 1

    def test_invalid_base64(self):
        from tools.excel_tools import read_excel
        with pytest.raises(Exception):
            read_excel("not_valid_base64!!!")

    def test_wrong_mime(self):
        from tools.excel_tools import read_excel
        # Pass a PDF magic header instead of xlsx
        fake_pdf = b64(b"%PDF-1.4 fake content")
        with pytest.raises(Exception):
            read_excel(fake_pdf)


# ─────────────────────────────────────────────────────────────────────────────
# Word / DOCX
# ─────────────────────────────────────────────────────────────────────────────

class TestTextToDocx:
    """text_to_docx — Markdown → .docx."""

    def test_basic_markdown(self):
        from tools.docx_tools import text_to_docx
        md = "# Hello World\n\nThis is a **bold** paragraph.\n\n- Item 1\n- Item 2"
        result = text_to_docx(md, title="Test Doc", author="LUKA")
        assert "file_content" in result
        raw = from_b64(result["file_content"])
        assert raw[:4] == b"PK\x03\x04"  # docx = zip

    def test_table_markdown(self):
        from tools.docx_tools import text_to_docx
        md = "| Name | Value |\n|------|-------|\n| Alice | 100 |"
        result = text_to_docx(md)
        assert result["size_bytes"] > 0

    def test_no_title(self):
        from tools.docx_tools import text_to_docx
        result = text_to_docx("Plain paragraph only.")
        assert from_b64(result["file_content"])[:4] == b"PK\x03\x04"

    def test_code_block(self):
        from tools.docx_tools import text_to_docx
        md = "```\nprint('hello')\n```"
        result = text_to_docx(md)
        assert result["size_bytes"] > 0


class TestReadDocx:
    """read_docx — parse .docx and return structured JSON."""

    def _make_docx(self, paragraphs: list[str], heading: str | None = None) -> str:
        from docx import Document
        doc = Document()
        if heading:
            doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
        buf = io.BytesIO()
        doc.save(buf)
        return b64(buf.getvalue())

    def test_basic_read(self):
        from tools.docx_tools import read_docx
        b64_file = self._make_docx(
            ["First paragraph.", "Second paragraph."],
            heading="Introduction",
        )
        result = read_docx(b64_file)
        assert result["section_count"] >= 3  # heading + 2 paragraphs
        assert "Introduction" in result["plain_text"]
        assert result["word_count"] > 0

    def test_empty_doc(self):
        from tools.docx_tools import read_docx
        b64_file = self._make_docx([])
        result = read_docx(b64_file)
        assert result["word_count"] == 0

    def test_wrong_mime(self):
        from tools.docx_tools import read_docx
        fake = b64(b"This is not a docx file at all")
        with pytest.raises(Exception):
            read_docx(fake)


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

class TestTextToPdf:
    """text_to_pdf — Markdown → PDF."""

    def test_basic(self):
        from tools.pdf_tools import text_to_pdf
        result = text_to_pdf("# Title\n\nHello world.", title="Test", author="LUKA")
        assert "file_content" in result
        raw = from_b64(result["file_content"])
        assert raw[:4] == b"%PDF"

    def test_all_elements(self):
        from tools.pdf_tools import text_to_pdf
        md = """# H1
## H2
### H3

Normal paragraph with **bold** and *italic*.

- Bullet one
- Bullet two

1. First
2. Second

| Col A | Col B |
|-------|-------|
| 1     | 2     |

> Blockquote text

---

```
code block
```
"""
        result = text_to_pdf(md, title="Full Test")
        assert from_b64(result["file_content"])[:4] == b"%PDF"

    def test_letter_page(self):
        from tools.pdf_tools import text_to_pdf
        result = text_to_pdf("Content", page_size="LETTER")
        assert result["mime_type"] == "application/pdf"

    def test_no_title(self):
        from tools.pdf_tools import text_to_pdf
        result = text_to_pdf("Just plain text without a title.")
        assert result["size_bytes"] > 0


class TestReadPdf:
    """read_pdf — parse PDF and return structured JSON."""

    def _make_pdf(self, text: str) -> str:
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import A4
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        doc.build([Paragraph(text, styles["Normal"])])
        return b64(buf.getvalue())

    def test_basic_read(self):
        from tools.pdf_tools import read_pdf
        b64_file = self._make_pdf("Hello from LUKA. This is a test PDF.")
        result = read_pdf(b64_file)
        assert result["total_pages"] >= 1
        assert result["word_count"] > 0
        assert "Hello" in result["plain_text"]

    def test_page_range(self):
        from tools.pdf_tools import read_pdf
        b64_file = self._make_pdf("Single page content.")
        result = read_pdf(b64_file, page_range="1")
        assert result["pages_read"] == 1

    def test_wrong_mime(self):
        from tools.pdf_tools import read_pdf
        fake = b64(b"PK\x03\x04 not a real pdf")
        with pytest.raises(Exception):
            read_pdf(fake)

    def test_metadata_present(self):
        from tools.pdf_tools import read_pdf
        b64_file = self._make_pdf("Metadata test.")
        result = read_pdf(b64_file)
        assert "metadata" in result
        assert "page_count" in result["metadata"]


# ─────────────────────────────────────────────────────────────────────────────
# API endpoint
# ─────────────────────────────────────────────────────────────────────────────

class TestInvokeEndpoint:
    """Test the /invoke REST endpoint via FastAPI TestClient."""

    @pytest.fixture(autouse=True)
    def client(self):
        from fastapi.testclient import TestClient
        from server import app
        self.client = TestClient(app)

    def test_health(self):
        r = self.client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"

    def test_list_tools(self):
        r = self.client.get("/tools")
        assert r.status_code == 200
        tools = r.json()["tools"]
        assert "read_excel" in tools
        assert "text_to_pdf" in tools

    def test_invoke_text_to_pdf(self):
        r = self.client.post("/invoke", json={
            "tool": "text_to_pdf",
            "input": {"content": "Hello world", "title": "API Test"},
        })
        assert r.status_code == 200
        body = r.json()
        assert body["tool"] == "text_to_pdf"
        assert "file_content" in body["result"]

    def test_invoke_text_to_docx(self):
        r = self.client.post("/invoke", json={
            "tool": "text_to_docx",
            "input": {"content": "# Heading\n\nBody text."},
        })
        assert r.status_code == 200
        assert "file_content" in r.json()["result"]

    def test_invoke_data_to_excel(self):
        r = self.client.post("/invoke", json={
            "tool": "data_to_excel",
            "input": {"data": [{"a": 1, "b": 2}, {"a": 3, "b": 4}]},
        })
        assert r.status_code == 200
        assert r.json()["result"]["row_count"] == 2

    def test_unknown_tool(self):
        r = self.client.post("/invoke", json={
            "tool": "nonexistent_tool",
            "input": {},
        })
        assert r.status_code == 404

    def test_invalid_params(self):
        r = self.client.post("/invoke", json={
            "tool": "text_to_pdf",
            "input": {},  # missing required 'content'
        })
        assert r.status_code == 422


# ─────────────────────────────────────────────────────────────────────────────
# File utils
# ─────────────────────────────────────────────────────────────────────────────

class TestFileUtils:
    def test_decode_b64_valid(self):
        from utils.file_utils import decode_b64
        original = b"hello world"
        encoded = base64.b64encode(original).decode()
        assert decode_b64(encoded) == original

    def test_decode_b64_invalid(self):
        from utils.file_utils import decode_b64
        with pytest.raises(ValueError):
            decode_b64("not!!valid!!base64###")

    def test_encode_b64(self):
        from utils.file_utils import encode_b64
        data = b"test data"
        result = encode_b64(data)
        assert base64.b64decode(result) == data

    def test_check_size_ok(self):
        from utils.file_utils import check_size
        check_size(b"small data")  # should not raise

    def test_check_size_too_large(self):
        from utils.file_utils import check_size, FileSizeError
        import config
        big = b"x" * (config.MAX_FILE_SIZE_BYTES + 1)
        with pytest.raises(FileSizeError):
            check_size(big)

    def test_temp_file_cleanup(self):
        from utils.file_utils import temp_file
        from config import TEMP_DIR
        data = b"temp test data"
        path_ref = None
        with temp_file(data, suffix=".bin") as path:
            path_ref = path
            assert path.exists()
        assert not path_ref.exists()
